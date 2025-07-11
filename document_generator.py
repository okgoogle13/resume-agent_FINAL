# document_generator.py
"""
Handles the core logic of generating documents as structured Markdown text.
This version includes improved export quality for DOCX and PDF and minor bug fixes.
"""
from api_clients import GeminiClient
from typing import List, Dict, Any
from docx import Document
from weasyprint import HTML, CSS
import io
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import markdown
import json

# Initialize the embedding model once when the module is loaded
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

def find_relevant_experiences(question: str, experiences: List[Dict[str, Any]], top_k: int = 3) -> List[Dict[str, Any]]:
    """Finds the most relevant career experiences based on a query string."""
    if not experiences or not question:
        return []
    
    # Create a single text representation for each experience
    experience_texts = [f"{exp.get('title', '')}. {exp.get('situation', '')} {exp.get('task', '')} {exp.get('action', '')} {exp.get('result', '')}" for exp in experiences]
    
    # Generate embeddings
    question_embedding = embedding_model.encode([question])
    experience_embeddings = embedding_model.encode(experience_texts)
    
    # Calculate cosine similarity
    similarities = cosine_similarity(question_embedding, experience_embeddings)[0]
    
    # Get the indices of the top-k most similar experiences
    top_indices = similarities.argsort()[-top_k:][::-1]
    
    return [experiences[i] for i in top_indices]

class DocumentGenerator:
    def __init__(self, gemini_client: GeminiClient):
        self.gemini_client = gemini_client

    def _generate_ai_content(self, prompt: str) -> str:
        """Helper function to call the Gemini API."""
        return self.gemini_client.generate_text(prompt)

    def _create_docx_from_markdown(self, markdown_content: str) -> bytes:
        """Creates a formatted DOCX from Markdown content."""
        doc = Document()
        for line in markdown_content.split('\n'):
            line = line.strip()
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_paragraph()
                p.add_run(line[4:]).bold = True
            elif line.startswith('- '):
                # Add paragraph with bullet style, ensuring text is not empty
                p_text = line[2:].strip()
                if p_text:
                    doc.add_paragraph(p_text, style='List Bullet')
            elif line:
                doc.add_paragraph(line)
        
        # Save document to a byte stream
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()

    def _get_theme_css(self, theme_name: str) -> str:
        """
        Returns CSS string for a given theme name, based on user's detailed theme document.
        Converts inches to cm (1 inch = 2.54 cm) and points for spacing where appropriate.
        """
        # Helper for margin conversion: 0.75" = 1.905cm, 0.8" = 2.032cm, 0.9" = 2.286cm
        # Spacing: 12pt between sections, 6pt between items

        if theme_name == "Professional Classic": # Matches "Theme 1: Professional Classic" from user doc
            return """
                @page {
                    size: A4;
                    margin: 1.9cm; /* Approx 0.75" */
                }
                body {
                    font-family: 'Times New Roman', Times, serif;
                    font-size: 11pt;
                    line-height: 1.4; /* Adjusted for TNR */
                    color: #2d3748; /* Charcoal for body text */
                    background-color: #ffffff; /* Pure white background */
                }
                h1 { /* Name */
                    font-family: 'Times New Roman', Times, serif;
                    font-size: 16pt;
                    font-weight: bold;
                    color: #1a365d; /* Deep Navy for name */
                    margin-top: 0;
                    margin-bottom: 0.2cm; /* Space before contact info */
                    text-align: left; /* As per implied single column structure */
                }
                h1 + p { /* Contact Info paragraph immediately after H1 (Name) */
                    font-family: 'Times New Roman', Times, serif;
                    font-size: 10pt;
                    color: #2d3748; /* Charcoal, or #718096 for dates/secondary per spec */
                    margin-top: 0;
                    margin-bottom: 0.8cm; /* Space after contact block */
                    line-height: 1.3;
                    text-align: left;
                }
                hr.name-divider { /* Thin horizontal line under name */
                    border: 0;
                    height: 0.5pt;
                    background-color: #1a365d; /* Navy */
                    margin-bottom: 0.8cm;
                }
                h2 { /* Section Headers */
                    font-family: 'Times New Roman', Times, serif;
                    font-size: 14pt;
                    font-weight: bold;
                    color: #1a365d; /* Deep Navy */
                    margin-top: 12pt; /* 12pt between sections */
                    margin-bottom: 8pt; /* Space before content in section */
                    border-bottom: 0.5pt solid #718096; /* Subtle gray divider lines */
                    padding-bottom: 2pt;
                }
                h3 { /* Subheadings (Job Titles, Company Names, Degree Titles) */
                    font-family: 'Times New Roman', Times, serif;
                    font-size: 12pt;
                    font-weight: bold;
                    color: #2d3748; /* Charcoal */
                    margin-top: 8pt; /* Space above job title */
                    margin-bottom: 2pt; /* Less space before company/date line */
                }
                h3 + p { /* Company | Location | Dates line OR Institution | Location | Year */
                    font-family: 'Times New Roman', Times, serif;
                    font-size: 10pt; /* Adjusted to match contact info spec */
                    font-style: normal; /* Dates/secondary info might not need italic if gray */
                    color: #718096; /* Subtle gray for dates and secondary info */
                    margin-top: 0;
                    margin-bottom: 6pt; /* 6pt between items */
                }
                ul {
                    list-style-type: disc;
                    padding-left: 20pt;
                    margin-top: 0;
                    margin-bottom: 6pt;
                }
                li {
                    margin-bottom: 4pt; /* Slightly less than 6pt for list items */
                    color: #2d3748; /* Body text color */
                }
                p { /* General paragraphs within sections */
                    margin-top: 0;
                    margin-bottom: 6pt; /* 6pt between items */
                    color: #2d3748; /* Body text color */
                }
                strong { font-weight: bold; } /* Ensure markdown bold is honored */
                em { font-style: italic; } /* Ensure markdown italic is honored */
                a { color: #1a365d; text-decoration: underline; }
            """
        elif theme_name == "Modern Minimalist": # Matches "Theme 2: Modern Minimalist" from user doc
            return """
                @page {
                    size: A4;
                    /* Margins: 0.8" top/bottom (2.032cm), 0.9" left/right (2.286cm) */
                    margin-top: 2cm; margin-bottom: 2cm;
                    margin-left: 2.3cm; margin-right: 2.3cm;
                    background-color: #fafafa; /* Off-white background */
                }
                body {
                    font-family: 'Helvetica', 'Arial', sans-serif;
                    font-size: 10.5pt;
                    line-height: 1.5;
                    color: #2d3748; /* Dark Gray for body text */
                }
                h1 { /* Name */
                    font-family: 'Helvetica', 'Arial', sans-serif;
                    font-size: 18pt;
                    font-weight: bold;
                    color: #4a5568; /* Slate Blue for name */
                    margin-top: 0;
                    margin-bottom: 0.1cm; /* Minimal space before contact */
                    text-align: left;
                }
                /* For asymmetrical header: Name left, contact right. This is hard with current markdown.
                   The current markdown has H1 then P for contact. We'll style H1+P for now.
                   A true asymmetrical would need different HTML/Markdown structure. */
                h1 + p { /* Contact Info */
                    font-family: 'Helvetica', 'Arial', sans-serif;
                    font-size: 9pt;
                    color: #2d3748; /* Dark Gray */
                    margin-top: 0;
                    margin-bottom: 18pt; /* Spacing between sections */
                    text-align: left; /* Default, will be single block */
                }
                /* Optional: thin accent line under contact info - not easily done with current h1+p */

                h2 { /* Section Headers */
                    font-family: 'Helvetica', 'Arial', sans-serif;
                    font-size: 13pt;
                    font-weight: bold; /* Spec says bold */
                    color: #4a5568; /* Slate Blue */
                    margin-top: 18pt; /* Spacing between sections */
                    margin-bottom: 10pt; /* Space before content */
                    padding-left: 8pt; /* For subtle left border space */
                    border-left: 2pt solid #38b2ac; /* Teal subtle left border */
                }
                h3 { /* Subheadings (Job Titles, Company Names) */
                    font-family: 'Helvetica', 'Arial', sans-serif;
                    font-size: 11pt;
                    /* font-weight: 600; /* semibold - CSS uses numeric or keywords like 'bold' */
                    font-weight: bold; /* Using bold as semibold might not be available */
                    color: #2d3748; /* Dark Gray */
                    margin-top: 10pt;
                    margin-bottom: 2pt;
                }
                h3 + p { /* Company | Location | Dates line */
                    font-family: 'Helvetica', 'Arial', sans-serif;
                    font-size: 9pt; /* Matched contact info size */
                    font-style: normal;
                    color: #4a5568; /* Slate Blue or Dark Gray */
                    margin-top: 0;
                    margin-bottom: 8pt; /* Spacing between items */
                }
                ul {
                    list-style-type: none; /* Remove default bullets */
                    padding-left: 15pt;
                    margin-top: 0;
                    margin-bottom: 8pt;
                }
                li {
                    margin-bottom: 5pt;
                    color: #2d3748; /* Body text color */
                    position: relative; /* For custom bullet positioning */
                    padding-left: 15pt; /* Space for custom bullet */
                }
                li::before { /* Circular bullet points in teal color */
                    content: '●'; /* Unicode circle character */
                    color: #38b2ac; /* Teal */
                    font-size: 10pt; /* Adjust size of bullet */
                    position: absolute;
                    left: 0;
                    top: 0.05em; /* Adjust vertical alignment */
                }
                p {
                    margin-top: 0;
                    margin-bottom: 8pt; /* Spacing between items */
                }
                strong { font-weight: bold; }
                em { font-style: italic; }
                a { color: #38b2ac; text-decoration: none; }
            """
        # Fallback to a very basic default if theme_name is unknown (or use Classic as default)
        # For now, my previous "Classic Professional" can act as a fallback if needed,
        # but the UI should only offer valid theme names.
        # Let's make the original "Classic Professional" the ultimate fallback.
        return """
            @page { size: A4; margin: 1.8cm; }
            body { font-family: 'Georgia', 'Times New Roman', serif; font-size: 11pt; line-height: 1.5; color: #333333; }
            h1 { font-family: 'Georgia', serif; font-size: 26pt; color: #222222; margin-bottom: 0.3cm; margin-top: 0; text-align: center; }
            h1 + p { font-family: 'Arial', 'Helvetica', sans-serif; font-size: 10pt; color: #454545; margin-top: 0; margin-bottom: 1cm; text-align: center; }
            h2 { font-family: 'Georgia', serif; font-size: 15pt; color: #333333; border-bottom: 1.5px solid #555555; padding-bottom: 3px; margin-top: 1cm; margin-bottom: 0.5cm; text-transform: uppercase; font-weight: bold;}
            h3 { font-family: 'Georgia', serif; font-size: 12pt; color: #333333; font-weight: bold; margin-top: 0.7cm; margin-bottom: 0.1cm;}
            h3 + p { font-family: 'Arial', 'Helvetica', sans-serif; font-size: 10pt; font-style: italic; color: #454545; margin-top:0; margin-bottom: 0.3cm;}
            ul { list-style-type: disc; padding-left: 20px; }
            li { margin-bottom: 0.25em; }
            p { margin-bottom: 0.4em; }
            a { color: #0000EE; text-decoration: underline; }
            strong { font-weight: bold; }
        """

    def _create_pdf_from_markdown(self, markdown_content: str, theme_name: str = "Professional Classic") -> bytes:
        """Creates a styled PDF from Markdown content using a specific theme."""
        html_content = markdown.markdown(markdown_content, extensions=['markdown.extensions.tables']) # Added tables extension
        
        # Get CSS based on theme_name
        theme_css_string = self._get_theme_css(theme_name)
        css = CSS(string=theme_css_string)
        
        # Base styles that might apply to all themes or are general structure
        # These are minimal now as most is in theme_css_string
        html_doc = HTML(string=html_content)
        return html_doc.write_pdf(stylesheets=[css])

    def generate_resume_markdown(self, user_profile: Dict, experiences: List[Dict[str, Any]]) -> str:
        """
        Generates the resume content as a Markdown string based on a predefined template
        and user data.
        """
        # Helper to get data or return empty string
        def get_data(data_dict, key, default=''):
            return data_dict.get(key, default) if data_dict and data_dict.get(key) else default

        # --- Main Information ---
        full_name = get_data(user_profile, 'full_name', 'Your Name')
        # resume_headline = get_data(user_profile, 'resume_headline', '') # Placeholder for a future field
        phone = get_data(user_profile, 'phone')
        email = get_data(user_profile, 'email')
        # location = get_data(user_profile, 'address') # Using 'address' as 'location'
        linkedin = get_data(user_profile, 'linkedin_url')
        contact_parts = [f"Telephone: {phone}" if phone else None,
                         f"Email: {email}" if email else None,
                        #  f"Address: {location}" if location else None, # Decided to use LinkedIn for brevity here
                         f"LinkedIn: {linkedin}" if linkedin else None]
        contact_line = " | ".join(filter(None, contact_parts))

        md_content = f"# {full_name}\n"
        # if resume_headline: # Add if field exists
        #     md_content += f"_{resume_headline}_\n" # Example: Italicized headline
        md_content += f"{contact_line}\n"

        # --- Career Summary ---
        career_summary = get_data(user_profile, 'professional_summary')
        if career_summary:
            md_content += "\n## CAREER SUMMARY\n"
            md_content += f"{career_summary}\n"

        # --- Education ---
        # Assuming education is stored in user_profile or a separate list.
        # For now, let's assume it's not explicitly structured in the DB like the template.
        # This section would need to be adapted if education data is available.
        # Example structure if education was a list of dicts in user_profile:
        # education_entries = get_data(user_profile, 'education', [])
        # if education_entries:
        #     md_content += "\n## EDUCATION\n"
        #     for edu in education_entries:
        #         md_content += f"**{get_data(edu, 'degree_title')}**\n"
        #         md_content += f"{get_data(edu, 'institution')} | {get_data(edu, 'location')} | {get_data(edu, 'graduation_year')}\n\n"
        # For now, this section is omitted as data structure is not present in user_profile for multiple degrees.
        # A single "education" field in user_profile could be added if needed.

        # --- Skills ---
        # Consolidating all 'related_skills' from experiences into one list for now.
        # The template's categorized skills are not directly supported by current DB.
        all_skills = []
        for exp in experiences:
            skills_str = get_data(exp, 'related_skills')
            if skills_str:
                all_skills.extend([s.strip() for s in skills_str.split(',') if s.strip()])

        unique_skills = sorted(list(set(all_skills))) # Get unique skills and sort them
        if unique_skills:
            md_content += "\n## SKILLS\n"
            # Display skills as a flat list for now
            for skill in unique_skills:
                md_content += f"- {skill}\n"
            md_content += "\n"


        # --- Professional Experience ---
        if experiences:
            md_content += "## PROFESSIONAL EXPERIENCE\n"
            for exp in experiences:
                md_content += f"\n### {get_data(exp, 'title')}\n"
                company_line = f"**{get_data(exp, 'company')}**"
                if get_data(exp, 'dates'): # Assuming location is not stored per experience
                    company_line += f" | {get_data(exp, 'dates')}"
                md_content += f"{company_line}\n"

                # Using 'resume_bullets' for responsibilities/achievements
                # The template has "Key Responsibilities" and "Key Achievement" which is more structured.
                # We'll list all resume_bullets under the job.
                # Future: Could try to parse resume_bullets if they follow a pattern, or add new DB fields.
                resume_bullets = get_data(exp, 'resume_bullets')
                if resume_bullets:
                    # md_content += "Key Responsibilities:\n" # Could add this if formatting all bullets as responsibilities
                    for bullet in resume_bullets.split('\n'):
                        if bullet.strip():
                            md_content += f"- {bullet.strip().lstrip('- ')}\n"
                md_content += "\n"
        
        # --- Certifications & Professional Development ---
        # This data is not currently in the user_profile or experiences.
        # This section would need new DB fields and UI elements to manage.
        # For now, this section is omitted.
        # Example if data was available:
        # certifications = get_data(user_profile, 'certifications', [])
        # if certifications:
        #     md_content += "\n## CERTIFICATIONS & PROFESSIONAL DEVELOPMENT\n"
        #     for cert in certifications:
        #         md_content += f"{get_data(cert, 'name')} | {get_data(cert, 'issuing_body')} | {get_data(cert, 'date')}\n"

        return md_content.strip()

    def generate_ksc_response(self, ksc_question: str, user_profile: Dict, experiences: List[Dict[str, Any]], company_intel: Dict[str, str], role_title: str) -> Dict[str, Any]:
        """Generates a KSC response and returns a dictionary with content."""
        relevant_experiences = find_relevant_experiences(ksc_question, experiences)
        experience_text = "\n\n".join([f"Title: {exp['title']}\nSituation: {exp['situation']}\nTask: {exp['task']}\nAction: {exp['action']}\nResult: {exp['result']}" for exp in relevant_experiences])
        
        prompt = f"""
        **Persona:** You are an expert career coach for the Australian Community Services sector. Your tone is professional and authentic, mirroring the user's personal style if provided.
        **User's Personal Style Profile:** {user_profile.get('style_profile', 'N/A')}
        **Company Intelligence:** {company_intel.get('values_mission', 'N/A')}
        **Task:** Write a compelling KSC response to the question below.
        **KSC Question:** "{ksc_question}"
        **Reasoning Framework:**
        1.  **Deconstruct:** Identify the core competency in the KSC question.
        2.  **Select Evidence:** Choose the strongest parts of the provided STAR stories to prove this competency.
        3.  **Draft:** Structure the response using the STAR method. Weave in the user's personal style and align with the company's values.
        **Candidate's Most Relevant Career Examples:**
        ---
        {experience_text if experience_text else "No specific examples provided."}
        ---
        **Output Format:** Generate clean, professional Markdown. Start directly with the response, do not add extra headings like "KSC Response".
        """
        markdown_content = self._generate_ai_content(prompt)
        return {"html": markdown_content}

    def generate_cover_letter_markdown(self, user_profile: Dict, experiences: List[Dict[str, Any]], job_details: Dict[str, Any], company_intel: Dict[str, Any]) -> str:
        """Generates a cover letter as a Markdown string."""
        job_desc_for_search = job_details.get('full_text', '')
        most_relevant_experience = find_relevant_experiences(job_desc_for_search, experiences, top_k=1)
        experience_snippet = ""
        if most_relevant_experience:
            exp = most_relevant_experience[0]
            experience_snippet = f"For instance, in my role as a {exp['title']} at {exp['company']}, I was responsible for {exp['task']}. I successfully {exp['action']}, which directly resulted in {exp['result']}."
        
        prompt = f"""
        **Persona:** You are an expert career advisor writing a cover letter for the Australian Community Services sector. Your tone is professional and warm, mirroring the user's personal style if provided.
        **User's Personal Style Profile:** {user_profile.get('style_profile', 'N/A')}
        **Company Intelligence:** {company_intel.get('values_mission', 'I am deeply impressed by your commitment to the community.')}
        **Task:** Write a compelling three-paragraph cover letter.
        **Reasoning Framework:**
        1.  **Opening:** State the role you are applying for and express genuine enthusiasm for the company, referencing a specific piece of company intelligence.
        2.  **Body:** Connect your skills directly to the key requirements of the job. Integrate the "Most Relevant Career Example" to provide concrete proof of your abilities and achievements.
        3.  **Closing:** Reiterate your strong interest in the role and the company. Include a clear call to action, stating your eagerness to discuss your application further.
        **Most Relevant Career Example:**
        ---
        {experience_snippet if experience_snippet else "The applicant has extensive experience directly relevant to this role's requirements."}
        ---
        **Output Format:** Generate clean, professional Markdown for a cover letter.
        """
        return self._generate_ai_content(prompt)

    def score_resume(self, resume_text: str, job_description: str) -> Dict[str, Any]:
        """Scores a resume against a job description using AI."""
        prompt = f"""
        **Persona:** You are an expert ATS (Applicant Tracking System) and a senior recruiter for the Community Services sector.
        **Task:** Analyze the provided resume against the job description. Provide a match score and actionable feedback.
        **Reasoning Framework:**
        1.  **Keyword Analysis:** Extract key skills, qualifications, and duties from the job description.
        2.  **Resume Parsing:** Identify skills, experiences, and achievements in the resume.
        3.  **Alignment Scoring:** Calculate a percentage score based on how well the resume matches the key requirements. Score harshly and realistically. A perfect match is rare.
        4.  **Feedback Generation:** Provide a list of strengths (what matched well) and a list of concrete, actionable suggestions for improvement (e.g., "Add the keyword 'Child Safety Framework' to your skills section," or "Quantify the achievement in your role at Hope Services by mentioning the number of clients served.").
        
        **Job Description:**
        ---
        {job_description}
        ---
        
        **Candidate's Resume:**
        ---
        {resume_text}
        ---
        
        **Output Format:**
        Return a single, valid JSON object only. Do not include any other text.
        {{
          "match_score": <integer_percentage>,
          "strengths": ["...", "..."],
          "suggestions": ["...", "..."]
        }}
        """
        response_text = self._generate_ai_content(prompt)
        try:
            # Clean up potential markdown formatting around the JSON
            json_str = response_text.strip().replace("```json", "").replace("```", "")
            return json.loads(json_str)
        except json.JSONDecodeError:
            return {"error": "Could not parse AI response.", "raw_text": response_text}

    def analyze_document_for_keywords_ats(self, resume_text: str, job_description_text: str) -> Dict[str, Any]:
        """
        Analyzes a resume against a job description for keyword matching and ATS friendliness.
        Orchestrates two AI calls:
        1. Extract critical keywords from the job description.
        2. Analyze the resume against these keywords and for ATS compatibility.
        """
        # --- Interaction 1: Job Description Keyword Extraction ---
        jd_prompt = f"""
Analyze the following job description. Extract key skills, technical terms, qualifications, and responsibilities.
Identify up to 10-15 critical keywords/phrases that are essential for an ATS and human reviewer.
For each keyword/phrase, categorize it as 'technical_skill', 'soft_skill', 'qualification', or 'responsibility'.
Provide the output as a single, valid JSON object only, with a key "critical_keywords", which is an array of objects, each containing "term" and "category".

Job Description:
---
{job_description_text}
---

Example Output:
{{
  "critical_keywords": [
    {{"term": "Project Management", "category": "technical_skill"}},
    {{"term": "Agile Methodology", "category": "technical_skill"}},
    {{"term": "Team Leadership", "category": "soft_skill"}}
  ]
}}
"""
        try:
            jd_response_text = self._generate_ai_content(jd_prompt)
            # Clean up potential markdown formatting
            jd_json_str = jd_response_text.strip().replace("```json", "").replace("```", "")
            jd_data = json.loads(jd_json_str)
            critical_keywords = jd_data.get("critical_keywords", [])
            if not critical_keywords: # Basic check if keywords were extracted
                 return {"error": "Could not extract critical keywords from job description.", "raw_jd_response": jd_response_text}
        except json.JSONDecodeError:
            return {"error": "Could not parse keyword extraction response from AI.", "raw_jd_response": jd_response_text}
        except Exception as e: # Catch other potential errors during keyword extraction
            return {"error": f"An unexpected error occurred during keyword extraction: {str(e)}", "raw_jd_response": jd_response_text if 'jd_response_text' in locals() else "N/A"}

        # --- Interaction 2: Resume Analysis against Keywords and ATS Friendliness ---
        resume_analysis_prompt = f"""
You are an expert ATS and resume analyst.
Analyze the provided resume based on the following critical keywords:
{json.dumps(critical_keywords)}

The resume text is:
---
{resume_text}
---

Provide your analysis as a single, valid JSON object only, with the following structure:
1. "keyword_analysis": An array of objects. For each critical keyword, indicate if it's "present" (true/false) in the resume. If "false", provide one brief, actionable "suggestion" on how to naturally integrate it. The "term" and "category" for each keyword should be included from the input.
2. "ats_friendliness": An object containing:
    - "overall_score": A score from 0-100 indicating ATS friendliness (0 poor, 100 excellent).
    - "issues": An array of objects, where each object has:
        - "issue_type": e.g., "Use of Tables", "Columns Detected", "Non-Standard Font", "Image Detected", "Complex Headers/Footers", "Lack of Keywords".
        - "description": A brief explanation of the issue.
        - "suggestion": A brief suggestion to fix it.
3. "general_suggestions": An array of 2-3 general actionable suggestions for improving the resume's keyword optimization or ATS compatibility.

Example Output:
{{
  "keyword_analysis": [
    {{"term": "Project Management", "category": "technical_skill", "present": true, "suggestion": null}},
    {{"term": "Agile Methodology", "category": "technical_skill", "present": false, "suggestion": "Consider adding 'Agile Methodology' to your project description where you managed the software development lifecycle."}}
  ],
  "ats_friendliness": {{
    "overall_score": 75,
    "issues": [
      {{"issue_type": "Columns Detected", "description": "The resume uses a two-column layout which might confuse some ATS.", "suggestion": "Consider a single-column layout for better ATS parsing."}}
    ]
  }},
  "general_suggestions": [
    "Ensure your contact information is at the top and easily parsable."
  ]
}}
"""
        try:
            resume_response_text = self._generate_ai_content(resume_analysis_prompt)
            # Clean up potential markdown formatting
            resume_json_str = resume_response_text.strip().replace("```json", "").replace("```", "")
            # The final response is the resume analysis data, but we'll also add the extracted keywords for the UI
            final_data = json.loads(resume_json_str)
            final_data["extracted_critical_keywords"] = critical_keywords # Add this for easier UI rendering
            return final_data
        except json.JSONDecodeError:
            return {"error": "Could not parse resume analysis response from AI.", "raw_resume_response": resume_response_text, "extracted_critical_keywords": critical_keywords}
        except Exception as e: # Catch other potential errors during resume analysis
            return {"error": f"An unexpected error occurred during resume analysis: {str(e)}", "raw_resume_response": resume_response_text if 'resume_response_text' in locals() else "N/A", "extracted_critical_keywords": critical_keywords}
